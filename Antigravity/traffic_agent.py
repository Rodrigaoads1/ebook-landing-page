import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import os
from datetime import datetime

class TrafficAgent:
    def __init__(self, brand_name="Migração Digital"):
        self.brand_name = brand_name
        # Caminho onde eu tenho permissão de execução
        self.base_path = r"c:\Users\Rodrigo\Downloads\ebook\Antigravity"
        self.data_folder = os.path.join(self.base_path, "dados_clientes")
        self.output_folder = os.path.join(self.base_path, "relatorios_gerados")
        
        # Garantir que as pastas existam
        for folder in [self.data_folder, self.output_folder]:
            if not os.path.exists(folder):
                os.makedirs(folder)

    def process_csv(self, file_path):
        """Lê o CSV tratando formatos e encodings diferentes."""
        encodings = ['utf-8', 'latin-1', 'utf-16', 'cp1252']
        df = None
        
        for enc in encodings:
            try:
                # Tenta ler a primeira linha para detectar Google Ads
                with open(file_path, 'r', encoding=enc) as f:
                    first_line = f.readline()
                    is_google_ads = "Relat" in first_line or " de " in first_line
                
                if is_google_ads:
                    df = pd.read_csv(file_path, skiprows=2, encoding=enc, sep=',')
                else:
                    df = pd.read_csv(file_path, encoding=enc)
                print(f"Sucesso ao ler {file_path} com encoding {enc}")
                break
            except Exception:
                continue
        
        if df is None:
            print(f"Erro: Não foi possível ler {file_path} com nenhum dos encodings testados.")
            return None

        try:
            # Limpeza de dados: remover linha de 'Total'
            if 'Campanha' in df.columns:
                df = df[df['Campanha'].astype(str).str.contains('Total', case=False) == False]
                df = df.dropna(subset=['Campanha'])

            # Mapeamento dinâmico: procura palavras-chave nas colunas
            final_mapping = {}
            keywords = {
                'Campaign': ['Campanha', 'Campaign'],
                'Clicks': ['Cliques', 'Clicks'],
                'Impressions': ['Impr', 'Impressions'],
                'Cost': ['Custo', 'Spend', 'Gasto'],
                'Leads': ['Todas as conv', 'Conve', 'Leads', 'Resultados'],
                'CTR': ['CTR']
            }
            
            # Primeiro tenta match EXATO para evitar confusão (ex: Custo vs Custo/Conv)
            unmapped_cols = list(df.columns)
            for standard_name, keys in keywords.items():
                # Tenta match exato primeiro
                for col in unmapped_cols:
                    if any(col.lower() == k.lower() for k in keys):
                        final_mapping[col] = standard_name
                        unmapped_cols.remove(col)
                        break
            
            # Depois tenta match parcial para o que sobrar
            for standard_name, keys in keywords.items():
                if standard_name in final_mapping.values(): continue
                for col in unmapped_cols:
                    if any(k.lower() in col.lower() for k in keys):
                        final_mapping[col] = standard_name
                        unmapped_cols.remove(col)
                        break
            
            df = df.rename(columns=final_mapping)
            print(f"Colunas traduzidas: {df.columns.tolist()}")

            # Converter colunas numéricas
            target_cols = ['Impressions', 'Clicks', 'Cost', 'Leads', 'CTR']
            for col in target_cols:
                if col in df.columns:
                    val = df[col].astype(str).str.replace('"', '').str.replace('%', '')
                    if ',' in val.iloc[0] if not val.empty else False:
                        val = val.str.replace('.', '', regex=False).str.replace(',', '.', regex=False)
                    df[col] = pd.to_numeric(val, errors='coerce').fillna(0)
                else:
                    # Cria a coluna com zeros se não existir para não quebrar o script
                    df[col] = 0

            return df
        except Exception as e:
            print(f"Erro ao processar dados de {file_path}: {e}")
            return None

    def generate_analysis(self, df):
        """Gera insights automáticos para o expert revisar."""
        insights = []
        if df is None or df.empty:
            return ["Dados insuficientes para análise."]

        try:
            total_cost = df['Cost'].sum()
            total_leads = df['Leads'].sum()
            avg_cpl = total_cost / total_leads if total_leads > 0 else 0
            
            insights.append(f"INVESTIMENTO TOTAL: R$ {total_cost:,.2f}")
            insights.append(f"LEADS GERADOS: {total_leads:.0f}")
            if avg_cpl > 0:
                insights.append(f"CPL MÉDIO: R$ {avg_cpl:,.2f}")

            # Identificar melhor campanha
            leads_df = df[df['Leads'] > 0]
            if not leads_df.empty:
                leads_df = leads_df.copy()
                leads_df['temp_cpl'] = leads_df['Cost'] / leads_df['Leads']
                top_campaign = leads_df.sort_values('temp_cpl').iloc[0]
                insights.append(f"DESTAQUE: A campanha '{top_campaign['Campaign']}' é a mais eficiente com CPL de R$ {top_campaign['temp_cpl']:,.2f}.")
        except Exception as e:
            insights.append(f"Aviso: Erro ao calcular insights detalhados ({e})")

        insights.append("REVISÃO NECESSÁRIA: [Adicione aqui seu comentário estratégico final para o cliente]")
        return insights

    def generate_charts(self, df, client_name):
        """Cria gráficos usando os dados REAIS do dataframe."""
        if df is None or df.empty or 'Cost' not in df.columns:
            return None

        plt.style.use('bmh')
        fig, ax1 = plt.subplots(figsize=(10, 5))

        color_clicks = '#003366' 
        color_impr = '#CCCCCC' 
        
        df_plot = df.sort_values(by='Cost', ascending=False).head(5)
        campaigns = df_plot['Campaign'].astype(str).str.slice(0, 20) + "..."
        x = range(len(campaigns))

        ax1.bar(x, df_plot['Impressions'], width=0.4, label='Impressões', color=color_impr, align='center')
        ax1.set_ylabel('Impressões', color='#666666')
        
        ax2 = ax1.twinx()
        ax2.plot(x, df_plot['Clicks'], marker='o', label='Cliques', color=color_clicks, linewidth=2)
        ax2.set_ylabel('Cliques', color=color_clicks)

        plt.title(f"Performance por Campanha: {client_name}", pad=20, fontsize=14, fontweight='bold')
        ax1.set_xticks(x)
        ax1.set_xticklabels(campaigns, rotation=15, ha='right')
        
        lines1, labels1 = ax1.get_legend_handles_labels()
        lines2, labels2 = ax2.get_legend_handles_labels()
        ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper right')

        plt.tight_layout()
        
        chart_path = os.path.join(self.output_folder, f"performance_{client_name}.png")
        plt.savefig(chart_path, dpi=150)
        plt.close()
        return chart_path

    def create_report(self, client_name, df, insights):
        doc = Document()
        
        # Cabeçalho "Logo" estilizado com texto
        section = doc.sections[0]
        header = section.header
        htable = header.add_table(1, 2, width=Inches(6))
        
        cell_logo = htable.cell(0, 0)
        p_logo = cell_logo.paragraphs[0]
        run_logo = p_logo.add_run("MIGRAÇÃO DIGITAL")
        run_logo.bold = True
        run_logo.font.size = Pt(18)
        run_logo.font.color.rgb = RGBColor(0, 51, 102) # Azul Escuro
        
        cell_info = htable.cell(0, 1)
        p_info = cell_info.paragraphs[0]
        p_info.text = f"Relatório Estratégico\n{datetime.now().strftime('%d/%m/%Y')}"
        p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Título Principal
        doc.add_heading(f"Performance de Tráfego: {client_name}", 0)
        
        # 1. Resumo Executivo
        doc.add_heading("1. Resumo Executivo", level=1)
        for insight in insights:
            p = doc.add_paragraph(insight, style='List Bullet')
            if "REVISÃO NECESSÁRIA" in insight:
                run = p.runs[0]
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True

        # 2. Análise Visual de Ativos
        doc.add_heading("2. Análise Visual", level=1)
        chart_path = self.generate_charts(df, client_name)
        if chart_path:
            doc.add_picture(chart_path, width=Inches(6))
            doc.add_paragraph("Gráfico: Comparativo de Impressões (cinza) e Cliques (azul).", style='Caption')

        # 3. Tabela de Detalhes
        doc.add_heading("3. Detalhamento", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Campanha'
        hdr_cells[1].text = 'Impr.'
        hdr_cells[2].text = 'Cliques'
        hdr_cells[3].text = 'Custo (R$)'
        hdr_cells[4].text = 'Leads'

        for _, row in df.head(10).iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Campaign'])[:30]
            row_cells[1].text = f"{row['Impressions']:,.0f}"
            row_cells[2].text = f"{row['Clicks']:,.0f}"
            row_cells[3].text = f"{row['Cost']:,.2f}"
            row_cells[4].text = f"{row['Leads']:,.0f}"

        output_file_name = f"Relatorio_{client_name}.docx"
        # Limpar nome do arquivo para evitar erros no Windows
        output_path = os.path.join(self.output_folder, output_file_name.replace(" ", "_"))
        doc.save(output_path)
        print(f"Relatório finalizado em: {output_path}")
        return output_path

    def process_all(self):
        """Processa todos os arquivos CSV na pasta de dados."""
        if not os.path.exists(self.data_folder):
            print(f"Pasta {self.data_folder} não encontrada.")
            return

        files = [f for f in os.listdir(self.data_folder) if f.endswith('.csv')]
        if not files:
            print("Nenhum arquivo CSV encontrado em 'dados_clientes'.")
            return

        for file_name in files:
            print(f"Processando {file_name}...")
            client_name = file_name.replace(".csv", "").replace("_", " ").title()
            file_path = os.path.join(self.data_folder, file_name)
            
            df = self.process_csv(file_path)
            if df is not None:
                insights = self.generate_analysis(df)
                self.create_report(client_name, df, insights)

if __name__ == "__main__":
    agent = TrafficAgent()
    print(f"Agente Ativado!")
    print(f"Localização Base: {agent.base_path}")
    print("-" * 30)
    agent.process_all()
    print("-" * 30)
    print("Processamento concluído.")
